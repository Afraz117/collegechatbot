import os
from typing import List
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

class DocumentParser:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
        except Exception as e:
            print(f"Error parsing PDF file {file_path}: {e}")
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error parsing DOCX file {file_path}: {e}")
        return text

    def extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"Error parsing TXT file {file_path}: {e}")
            return ""

    def parse_file(self, file_path: str) -> List[Document]:
        _, ext = os.path.splitext(file_path.lower())
        text = ""

        if ext == ".pdf":
            text = self.extract_text_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            text = self.extract_text_from_docx(file_path)
        elif ext == ".txt":
            text = self.extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file format: {ext}")
            return []

        if not text.strip():
            return []

        # Split the text into document chunks
        filename = os.path.basename(file_path)
        chunks = self.text_splitter.split_text(text)
        
        documents = [
            Document(
                page_content=chunk,
                metadata={"source": filename, "path": file_path}
            )
            for chunk in chunks
        ]
        
        return documents
